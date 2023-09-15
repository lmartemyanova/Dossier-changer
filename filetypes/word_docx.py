from win32com import client as wc
import os
import docx


class WordDocx:

    def __init__(self, file):
        """
        to initialize the class object
        :param file: the filename of the file
        """

        self.file = file

    def rename(self, filenames_attributes: dict, path: str) -> None:
        """
        To change the filename:
        If filename has specific text markers
        (as name of the manufacturer etc).
        Markers should be placed to dict(file_names)
        as "old marker": "new marker"
        :param filenames_attributes: the dict from user input (key = phrase to change, value = phrase for change)
        :param path: the path to folder
        :return: None
        """

        for name in filenames_attributes:
            if name in self.file and self.file.endswith('docx') and not self.file.startswith('~'):
                old_name = os.path.join(path, self.file)
                n = os.path.splitext(self.file)[0]
                b = n.split('_', maxsplit=2)
                new = b[0] + '_' + b[1] + '_' + filenames_attributes[name] + '.docx'
                new_name = os.path.join(path, new)
                os.rename(old_name, new_name)
        return

    def replace_text(self, text_dict: dict, path: str) -> None:
        """
        To replace text, containing in param text_dict (key to value).
        :param text_dict: the text to be replaced
        :param path: the path to folder
        :return: None
        """

        if self.file.endswith('docx') and not self.file.startswith('~'):
            doc = docx.Document(path)
            style = doc.styles['Normal']
            font = style.font
            for i in text_dict:
                for paragraph in doc.paragraphs:
                    if paragraph.text.find(i) >= 0:
                        paragraph.text = paragraph.text.replace(i, text_dict[i])
            for j in text_dict:
                for table in doc.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph.text.find(j) >= 0:
                                    paragraph.text = paragraph.text.replace(j, text_dict[j])
            for k in text_dict:
                for section in doc.sections:
                    header = section.header
                    for paragraph in header.paragraphs:
                        if paragraph.text.find(k) >= 0:
                            paragraph.text = paragraph.text.replace(k, text_dict[k])
                            # style = doc.styles['Normal']
                            # font = style.font
                            font.name = 'Times New Roman'
                            font.size = docx.shared.Pt(9)
                    for table in header.tables:
                        for col in table.columns:
                            for cell in col.cells:
                                for paragraph in cell.paragraphs:
                                    if paragraph.text.find(k) >= 0:
                                        paragraph.text = paragraph.text.replace(k, text_dict[k])
                                        # style = doc.styles['Normal']
                                        # font = style.font
            doc.save(os.path.basename(path))
            return

    def find_usages(self, phrase: str, path: str) -> str:
        """
        The method to find usages of any phrase in many files from user input
        :param phrase: phrase from user
        :param path: the path to folder
        :return: the name of the files where the usages of the phrase were founded
                 in case of Exceptions, return the filename with the Exception.
        """

        if self.file.endswith('docx') and not self.file.startswith('~'):
            path = os.path.join(path, self.file)
            try:
                doc = docx.Document(path)
            except Exception as e:
                return f'Возникла ошибка {e} при чтении файла: {self.file}'
            for paragraph in doc.paragraphs:
                if paragraph.text.find(phrase) >= 0:
                    return self.file
            for table in doc.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.find(phrase) >= 0:
                                return self.file
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    if paragraph.text.find(phrase) >= 0:
                        return self.file
                for table in header.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph.text.find(phrase) >= 0:
                                    return self.file


class WordDoc(WordDocx):

    def save_docx(self, path: str) -> None:
        """
        To save a .doc file as .docx
        :param path: the path to folder
        :return: None
        """

        w = wc.Dispatch('Word.Application')
        # path = os.path.join(path, self.file)
        file = self.file
        if file.endswith('doc') and not file.startswith('~'):
            doc = w.Documents.Open(path)
            doc.SaveAs(path + "x", 16)
            doc.Close()
        w.Quit()
        return
