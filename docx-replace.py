# pip install python-docx
import docx
import os

Dictionary = {

             }

def rename():
    file_names = {
                 
                 }
    for name in file_names:
        for file in os.listdir(folder):
            if name in file and file.endswith('docx') and not file.startswith('~'):
                old_name = os.path.join(folder, file)
                n = os.path.splitext(file)[0]
                b = n.split('_', maxsplit = 2)
                new = b[0] + '_' + b[1] + '_' + file_names[name] + '.docx'
                new_name = os.path.join(folder, new)
                os.rename(old_name, new_name)

paths = []
folder = os.getcwd()
for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('docx') and not file.startswith('~'):
            paths.append(os.path.join(root, file))

for path in paths:
    doc = docx.Document(path)
    style = doc.styles['Normal']
    font = style.font

    for i in Dictionary:
        for paragraph in doc.paragraphs:
            if paragraph.text.find(i) >= 0:
                paragraph.text = paragraph.text.replace(i, Dictionary[i])

    for j in Dictionary:
        for table in doc.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.find(j) >= 0:
                            paragraph.text = paragraph.text.replace(j, Dictionary[j])

    for k in Dictionary:
        for section in doc.sections:
            header = section.header

            for paragraph in header.paragraphs:
                if paragraph.text.find(k) >= 0:
                    paragraph.text = paragraph.text.replace(k, Dictionary[k])
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Times New Roman'
                    font.size = docx.shared.Pt(9)

            for table in header.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.find(k) >= 0:
                                paragraph.text = paragraph.text.replace(k, Dictionary[k])
                                style = doc.styles['Normal']
                                font = style.font
                                font.name = 'Times New Roman'
                                font.size = docx.shared.Pt(9)

    doc.save(os.path.basename(path))
    rename()